import requests
from io import BytesIO
from PIL import Image
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.shared import RGBColor, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import _Cell

# import docx
from haggis.files.docx import list_number
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import traceback
import typing
from datetime import datetime
import os
import glob

# import lxml

# HTML Content
html_content = open("test.html", "r").read()  # Replace with the actual HTML content
# html_content = '<p>Lorem ipsum dolor sit amet,<span style="color:#33ff57">consectetur adipiscing elit</span>. Integer nec odio.</p><p>Praesent libero. Sed cursus ante dapibus diam.<span style="color:#ff33a8;background-color:#d3d3d3">Sed nisi</span>.</p><p id="abd">Nulla quis sem at nibh elementum imperdiet.<span style="color:#33fff5">Duis sagittis ipsum</span>.</p><p id="abd"> Nulla quis <strong><em> <u>sem at </u></em> sem nibh</strong>elementum<b> imperdiet.</b><span style="color:#11afd3"> Duis sagittis ipsum </span>.</p>'

# Parse HTML with BeautifulSoup
soup = BeautifulSoup(html_content, "lxml")

# Create a Word document
doc = Document()
doc.core_properties.title = soup.title.string if soup.title else "HTML to DOCX"
doc.add_heading(doc.core_properties.title, level=0)


# Function to convert CSS color to hex code
def extract_hex_color(style: str) -> dict[str, str]:
    color_dict = {}
    try:
        if "color:" in style:
            color = style.split("color:")[1].split(";")[0].strip()
            if (
                color.startswith("#") and len(color) == 7
            ):  # Ensure it's a valid hex code
                color_dict["color"] = color.lstrip("#")  # Remove the '#' for RGBColor
        if "background-color:" in style:
            color = style.split("background-color:")[1].split(";")[0].strip()
            if (
                color.startswith("#") and len(color) == 7
            ):  # Ensure it's a valid hex code
                color_dict["background-color"] = color.lstrip(
                    "#"
                )  # Remove the '#' for RGBColor
    except IndexError:
        pass
    return color_dict  # Return None if no valid color is found


# Function to add styled text
def add_styled_text(paragraph, text, color=None, is_bg_color=False):
    run = paragraph.add_run(text)
    if color:
        try:
            if "color" in color["span"]:
                text_color = color.get("span").get("color", "000000")
                rgb = tuple(int(text_color[i : i + 2], 16) for i in (0, 2, 4))
                run.font.color.rgb = RGBColor(*rgb)
            if "bold" in color:
                run.bold = color["bold"]
            if "italic" in color:
                run.italic = color["italic"]
            if "underline" in color:
                run.underline = color["underline"]
            # if is_bg_color:
            #     bg_color = color.get("background-color", "FFFFFF")
            #     # rgb = tuple(int(bg_color[i:i + 2], 16) for i in (0, 2, 4))
            #     highlight = parse_xml(
            #         r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), bg_color)
            #     )
            #     parent_element = run._element.getparent()
            #     parent_element.insert(parent_element.index(run._element) + 1, highlight)
        except ValueError:
            pass  # Skip if the color value is invalid


# Function to handle nested lists
def process_list(doc_obj: _Cell, doc: Document, list_tag, parent_paragraph=None, level=1):
    for li in list_tag.find_all("li", recursive=False):
        for child in li.children:
            if child.name != "ol" and child.name != "ul":
                document_root = doc_obj if doc_obj else doc
                paragraph = document_root.add_paragraph(
                    style="List Number" if list_tag.name == "ol" else "List Bullet"
                )
                paragraph.paragraph_format.left_indent = Inches(level * 0.25)
            if child.name == "a":  # Handle anchor tags
                add_styled_text(paragraph, child.text.strip())
            # elif child.name == "span":
            #     span_color = extract_hex_color(child.get("style", ""))
            #     add_styled_text(paragraph, child.text.strip(), span_color)
            elif child.name == "p":
                p_style = parse_styles(child)
                process_p_child_tags(paragraph, child, li, p_style)
                if paragraph.style.name == "List Number":
                    list_number(doc, paragraph, prev=parent_paragraph)
                parent_paragraph = paragraph
            elif child.name == "ol" or child.name == "ul":
                process_list(doc_obj, doc, child, parent_paragraph=None, level=level + 1)


# Function to process tables
def process_table(table_tag):

    def process_table_cell(cell: Tag, docx_cell: _Cell):
        for tag in cell.children:
            if tag.name == "p":
                paragraph = docx_cell.add_paragraph()
                p_style = parse_styles(tag)
                process_p_child_tags(paragraph, tag, cell, p_style)
            elif tag.name == "img":
                img_data = add_image(doc, tag)
                if isinstance(img_data, BytesIO):
                    docx_cell.add_paragraph().add_run().add_picture(img_data)
                elif isinstance(img_data, str):
                    docx_cell.add_paragraph(img_data)
            elif tag.name == "ul" or tag.name == "ol":
                process_list(docx_cell, doc, tag)
                docx_cell.width = Inches(8.0)
            elif isinstance(tag, str):
                docx_cell.add_paragraph(tag)

    rows = table_tag.find_all("tr")
    table = doc.add_table(rows=0, cols=len(rows[0].find_all(["th", "td"])))
    table.autofit = False
    table.style = "Table Grid"

    for row in rows:
        cells = row.find_all(["th", "td"])
        row_cells = table.add_row().cells
        for i, cell in enumerate(cells):
            process_table_cell(cell, row_cells[i])


# Function to add an image
def add_image(img_tag: Tag):
    if "src" in img_tag.attrs:
        img_url = img_tag["src"]
        try:
            response = requests.get(img_url)
            response.raise_for_status()
            image = Image.open(BytesIO(response.content))
            img_buffer = BytesIO()
            image.thumbnail((300, 300))  # Resize image
            image.save(img_buffer, format=image.format)
            img_buffer.seek(0)
            return img_buffer
        except Exception as e:
            traceback.print_exc()
            print(f"Failed to load image: {img_url}, Error: {e}")
            return "Image not found"


def add_hyperlink(paragraph, text, url):
    """
    A function that adds a hyperlink to a paragraph.
    """
    pass


def set_paragraph_id(paragraph: Paragraph, p_id: str) -> None:
    """Sets paragraph ID using correct Word XML namespace"""
    try:
        # Create paragraph properties element if it doesn't exist
        if not paragraph._element.pPr:
            paragraph._element.get_or_add_pPr()

        # Create paragraphId tag with proper namespace
        para_id = parse_xml(
            f'<w:paraId w:val="{p_id}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        paragraph._element.pPr.append(para_id)
    except Exception as e:
        print(f"Failed to set paragraph ID: {str(e)}")


def parse_styles(tag: Tag) -> dict[str, str]:
    """Parses styles from a tag and merges them with the existing styles"""
    style_d = {}
    if "style" in tag.attrs:
        style = tag["style"]
        for s in style.split(";"):
            key, value = s.split(":")
            style_d[key.strip()] = (
                value.strip().lstrip("#") if "#" in value else value.strip()
            )
    return style_d

def process_p_child_tags(
    paragraph: Paragraph, tag: Tag, parent_tag: Tag, styles=None
) -> None:
    """Processes child tags of a paragraph"""
    for child in tag.children:
        ancestors = [tag.name] + [t.name for t in tag.parents]
        if isinstance(child, str) and child.parent.name == tag.name:
            if not styles:
                run = paragraph.add_run()
                run.text = child
            else:
                if "strong" not in ancestors and "b" not in ancestors:
                    styles["bold"] = False
                if "em" not in ancestors and "i" not in ancestors:
                    styles["italic"] = False
                if "u" not in ancestors:
                    styles["underline"] = False
                if "span" not in ancestors:
                    styles["span"] = {}
                add_styled_text(paragraph, child, styles)
        elif child.name == "p":
            process_p_child_tags(paragraph, child, tag, styles)
        elif child.name == "span":
            span_styles = parse_styles(child)
            styles.update({"span": span_styles})
            process_p_child_tags(paragraph, child, tag, styles)
        elif child.name == "strong" or child.name == "b":
            bold_styles = {"bold": True}
            styles.update(bold_styles)
            if "span" not in ancestors:
                styles["span"] = {}
            process_p_child_tags(paragraph, child, tag, styles)
        elif child.name == "em" or child.name == "i":
            italic_styles = {"italic": True}
            styles.update(italic_styles)
            if "strong" not in ancestors and "b" not in ancestors:
                styles["bold"] = False
            if "span" not in ancestors:
                styles["span"] = {}
            process_p_child_tags(paragraph, child, tag, styles)
        elif child.name == "u":
            underline_styles = {"underline": True}
            styles.update(underline_styles)
            if "em" != tag.name:
                styles["italic"] = False
            if tag.name not in ancestors:
                styles["bold"] = False
            if "span" not in ancestors:
                styles["span"] = {}
            process_p_child_tags(paragraph, child, tag, styles)


# Process the HTML content
bg_color = False
for tag in soup.body.descendants:
    if tag.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
        h_level = int(tag.name[1])
        doc.add_heading(tag.text.strip(), level=h_level)
    elif tag.name == "p" and tag.parent.name != "li":
        paragraph = doc.add_paragraph()
        process_p_child_tags(paragraph, tag, tag, {})
        p_id = tag.get("id", "")
        """ if p_id:
            set_paragraph_id(paragraph, p_id)
        color = {}
        if "style" in tag.attrs:
            color = extract_hex_color(tag["style"])
        for child in tag.children:
            if child.name == "a":  # Handle anchor tags
                # add_styled_text(paragraph, child.text.strip(), color)
                add_hyperlink(paragraph, child.text.strip(), child["href"])
            elif child.name == "span":
                span_color = extract_hex_color(child.get("style", ""))
                bg_color = True if "background-color" in child["style"] else False
                add_styled_text(
                    paragraph, child.text.strip(), span_color, is_bg_color=bg_color
                )
            elif child.name == "img":  # Handle image tags
                add_image(doc, child)
            elif isinstance(child, str):
                add_styled_text(paragraph, child, color) """
    elif tag.name == "table":
        process_table(tag)
    elif (tag.name == "ol" or tag.name == "ul") and tag.parent.name not in ["li", "td", "th"]:
        process_list(None, doc, tag)
    elif tag.name == "img":
        img = add_image(tag)
        if isinstance(img, BytesIO):
            doc.add_picture(img)
            img.close()
        elif isinstance(img, str):
            doc.add_paragraph(img)

# Delete previous created .docx files
docx_files = glob.glob("output_*.docx")
for file in docx_files:
    try:
        os.remove(file)
        print(f"Deleted file: {file}")
    except Exception as e:
        print(f"Failed to delete file: {file}, Error: {e}")

# Save the document
timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
output_filename = f"output_{timestamp}.docx"
doc.save(output_filename)
print("Document created successfully!")
