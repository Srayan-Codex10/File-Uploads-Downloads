import traceback
import base64
import re
from io import BytesIO
from bs4 import BeautifulSoup, Tag
from PIL import Image
from docx import Document
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.shared import RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.shared import OxmlElement, qn
from haggis.files.docx import list_number

soup = None


def check_anchor_id_length(html: str, article_id: int) -> str:
    """
    Fetch anchor links <a> by "href" value and check the id length
    If length is greater than 40 characters, then extra characters
    from the end are trimmed.
    Returns the modified HTML as string
    """
    soup_html = BeautifulSoup(html, "lxml")

    # Regular expression to match anchor ids greather 40 characters
    anchor_href_pattern = r"#[a-zA-Z0-9-]{41,}"

    # Retrieve all <a> tags where href matches the regular expression and trim the id
    anchor_tags = soup_html.body.find_all("a", href=re.compile(anchor_href_pattern))
    for a in anchor_tags:
        href_val = a["href"].split("#")[1]
        anchor_target = soup_html.find(id=href_val)
        if anchor_target:
            anchor_target["id"] = f"{href_val[:40]}"
            a["href"] = f"#{href_val[:40]}"
    return str(soup_html)


def add_href_anchor_tags(html):
    """
    This Method parses html string and adds href to all anchor tags
    with 'onclick' or 'ng-click' attributes
    Custom method to handle 'jumptosection' onclick events - document specific
    """
    pattern_1 = r"(?<=jumptosection\(')([a-zA-Z0-9-]+)(?='\);)"
    parsed_html = BeautifulSoup(html, "lxml")
    # Iterate over anchor tags with attribtues
    for a_tag in parsed_html.find_all("a", attrs={"onclick": True}):
        if a_tag["onclick"]:
            match = match_pattern(a_tag["onclick"], pattern_1)
            if match:
                a_tag["href"] = f"#{match}"
    return str(parsed_html)


def match_pattern(text: str, pattern: str):
    """
    Uses a Regular Expression to extract ID from onclick attribute value
    """
    result = re.search(pattern, text)
    if result:
        return result.group()
    else:
        return None


def skip_crlf(content):
    """
    Skip Adding empty '\\n' or '\\r' in the content
    New lines will be added from <br> tag or existing CRLF tokens at the end of text.
    """
    if not (isinstance(content, str)):
        return False
    if content != "\n" and content != "\r":
        return False
    else:
        return True


def add_docx_tables(doc: Document, table_html: Tag):
    """
    Add a docx table object to the Document object from html table tag.
    It extracts all text content and style information from 'style' attributes
    """

    def process_table_cell(cell: Tag, docx_cell: _Cell):
        style_attrs = None if "style" not in cell.attrs else cell["style"]
        style_data = parse_styles(style_attrs)
        process_p_child_tags(
            doc, docx_cell.add_paragraph(), cell, cell, docx_cell, style_data
        )

    if not table_html:
        return

    rows = table_html.find_all("tr")
    columns = max(len(row.find_all(["td", "th"])) for row in rows)
    docx_table = doc.add_table(rows=len(rows), cols=columns, style="Table Grid")

    for i, row in enumerate(rows):
        cells = row.find_all(["td", "th"])
        for j, cell in enumerate(cells):
            table_cell = docx_table.cell(i, j)
            process_table_cell(cell, table_cell)

    return docx_table


def create_bookmark_run(paragraph: Paragraph, bookmark_name: str, text: str, id: str):
    """
    Insert text in `paragraph` and surround it with bookmarkStart and bookmarkEnd,
    effectively creating a bookmark within the paragraph.
    """
    # Create a new run for the text
    run = paragraph.add_run(text)

    # Get the run's XML element
    r = run._r

    # Generate a unique ID for the bookmark. In Word, bookmark IDs must be numeric.
    # You can maintain a global counter or dictionary to ensure uniqueness if needed.
    # bookmark_id = str(abs(hash(bookmark_name)) % (10**6))

    # --- bookmarkStart ---
    tag_bookmark_start = OxmlElement("w:bookmarkStart")
    tag_bookmark_start.set(qn("w:id"), id)
    tag_bookmark_start.set(qn("w:name"), bookmark_name)

    # --- bookmarkEnd ---
    tag_bookmark_end = OxmlElement("w:bookmarkEnd")
    tag_bookmark_end.set(qn("w:id"), id)

    # Insert them around the text run in the XML
    r.insert_element_before(tag_bookmark_start)
    r.append(tag_bookmark_end)


def create_internal_hyperlink_run(
    paragraph: Paragraph, display_text: str, bookmark_name: str
):
    """
    Insert a run in `paragraph` that links (anchors) to the given bookmark_name within the same document.
    """
    # Create the <w:hyperlink> element and specify the anchor (bookmark target)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    hyperlink.set(
        qn("w:history"), "1"
    )  # This just indicates Word should store the link history

    # Create a <w:r> node to hold the text
    new_run = Run(OxmlElement("w:r"), paragraph)
    # Create a <w:rPr> for run properties (e.g., formatting)
    r_pr = OxmlElement("w:rPr")

    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)

    # Create <w:t> element (text) inside the run
    w_t = OxmlElement("w:t")
    w_t.text = display_text

    new_run._element.append(r_pr)
    new_run._element.append(w_t)
    new_run.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    new_run.font.underline = True

    # Add this <w:r> run into the <w:hyperlink>
    hyperlink.append(new_run._element)

    # Append the hyperlink into the paragraph
    paragraph._p.append(hyperlink)


def add_links(paragraph: Paragraph, text: str, url: str):
    """
    Add Hyperlinks to document based on <a> tags in html
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(
        qn("r:id"),
        r_id,
    )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = Run(OxmlElement("w:r"), paragraph)
    new_run.text = text

    # new_run.font.underline = True
    new_run.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    new_run.font.underline = True

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)


def process_list(
    docx_cell: _Cell, doc: Document, list_tag: Tag, parent_paragraph=None, level=1
) -> Paragraph:
    """
    Convert HTML <ul> and <ol> tags containing <li> recursively into docx bullets
    """
    for li in list_tag.find_all("li", recursive=False):
        for child in li.children:
            if child.name != "ol" and child.name != "ul" and (not skip_crlf(child)):
                doc_obj = docx_cell if docx_cell else doc
                paragraph = doc_obj.add_paragraph(
                    style="List Number" if list_tag.name == "ol" else "List Bullet"
                )
                paragraph.paragraph_format.left_indent = Inches(level * 0.25)
            if child.name == "a":  # Handle anchor tags
                add_links(paragraph, child.text, child.get("href", ""))
            elif child.name in ["p", "blockquote"]:
                if paragraph.style.name == "List Number":
                    list_number(doc, paragraph, prev=parent_paragraph)
                style_str = None if "style" not in child.attrs else child["style"]
                style_dict = parse_styles(style_str)
                align_para(style_dict, paragraph)
                if "id" in child.attrs:
                    create_bookmark_run(paragraph, child["id"], "", child["id"])
                process_p_child_tags(
                    doc, paragraph, child, child, docx_cell, style_dict
                )
                parent_paragraph = paragraph
            elif child.name == "ol" or child.name == "ul":
                process_list(
                    docx_cell, doc, child, parent_paragraph=None, level=level + 1
                )
            elif child.name == "table" and child.parent.name == "li":
                table = add_docx_tables(doc, child)
                tbl, p = table._tbl, paragraph._p
                p.addnext(tbl)

    return parent_paragraph


def is_list_continued(list_tag: Tag):
    """
    Check if an <ol> list is flat with only one <li>
    and continuing sequence
    """
    next_tags = list_tag.find_next_siblings("ol")
    prev_tags = list_tag.find_previous_siblings("ol")

    # Check following "ol" tags
    if len(next_tags) == 0 and len(prev_tags) == 0:
        return False
    elif len(prev_tags) > 0 and len(next_tags) == 0:
        return True

    next_sibling = next_tags[0]
    # check <ol> with "start" and no "value" attribute in <li>
    if next_sibling.name == "ol":
        if "start" in next_sibling.attrs and "start" not in list_tag.attrs:
            return False
        elif "start" in next_sibling.attrs and "start" in list_tag.attrs:
            return True
        elif "start" not in next_sibling.attrs and "start" in list_tag.attrs:
            return True
    elif next_sibling.name != "ol" and next_sibling.name != "ul":
        return False
    # check li tags
    for li in list_tag.find_all("li", recursive=False):
        if "value" in li.attrs:
            current_li_val = li["value"]
            sibling_li = next_sibling.find_all("li", recursive=False)
            for sib_li in sibling_li:
                if "value" in sib_li.attrs:
                    return int(sib_li["value"]) - int(current_li_val) == 1
                else:
                    return False
        else:
            return False


def add_images(img: Tag):
    """
    Adds Picture to docx Document
    Downloads image from src url or converts base64 encoded data
    """
    if "src" in img.attrs:
        img_url = img["src"]
        err_msg_https = f"Image not available, {img_url}"
        err_msg_img_data = "Image not available"
        image = None
        try:
            # Download image
            if img_url.startswith("https"):
                response = requests.get(img_url)
                response.raise_for_status()
                image = Image.open(BytesIO(response.content))
            # Convert raw base64 encoded image data
            elif img_url.startswith("data"):
                img_b64 = img_url.split(",")
                data = base64.urlsafe_b64decode(img_b64[1])
                img_format = img_b64[0].split(";")[0].split("/")[1].upper()
                image = Image.open(BytesIO(data), formats=[img_format, "JPEG"])
            img_buff = BytesIO()
            image.thumbnail((500, 400), Image.Resampling.LANCZOS)
            image.save(img_buff, format=image.format, quality=90)
            img_buff.seek(0)
            return img_buff
        except Exception as e:
            err = err_msg_https if img_url.startswith("https") else err_msg_img_data
            traceback.print_exc()
            return err
    # if html <img> tag does not have src attribute
    else:
        return "Image url not found"


def rgb_to_hex(rgb: str):
    """
    Parses an rgb string like rbg(233,42,12) to return #E92A0C
    """
    match = re.match(r"rgb\((\d{1,3}),\s*(\d{1,3}),\s*(\d{1,3})\)", rgb)
    if match:
        return "{:02X}{:02X}{:02X}".format(*map(int, match.groups()))
    return rgb


def parse_styles(style_attr: str) -> dict[str, str]:
    """
    Create a python dictionary of styles from html style attribute string
    Ex: style=color: red;background-color: blue
    Returns: {color: red, background-color: blue}
    """
    style_dict = {}
    if not style_attr:
        return style_dict
    for style_itm in style_attr.split(";"):
        if ":" in style_itm:
            k, v = style_itm.split(":")
            style_dict[k.strip()] = (
                v.strip().lstrip("#") if "#" in v else rgb_to_hex(v.strip())
            )
    return style_dict


def add_text_color(paragraph, text, styles=None, is_bg_color=False) -> None:
    """
    Add Styles to paragraph based on given style attributes from html
    """
    run = paragraph.add_run(text)
    if styles:
        try:
            if "color" in styles.get("span", ""):
                text_color = styles.get("span").get("color", "000000")
                rgb = tuple(int(text_color[i : i + 2], 16) for i in (0, 2, 4))
                run.font.color.rgb = RGBColor(*rgb)
            if "bold" in styles:
                run.bold = styles["bold"]
            if "italic" in styles:
                run.italic = styles["italic"]
            if "underline" in styles:
                run.underline = styles["underline"]
            if "font-family" in styles and "bold" in styles.get("font-family"):
                run.bold = True
        except ValueError:
            pass  # Skip if the color value is invalid


def tuple_check(tag, style, ancestors, style_d):
    tag_1, tag_2 = tag
    if (tag_1 not in ancestors) and (tag_2 not in ancestors):
        style_d[style] = False
    else:
        style_d[style] = True


def check_style_parent(ancestors: [], style_d: dict, tag_names: dict) -> dict[str, str]:
    """
    Remove the styles which are not in any direct parent or ancestor of the current tag
    """
    for tag, style in tag_names.items():
        if isinstance(tag, tuple):
            tuple_check(tag, style, ancestors, style_d)
        else:
            if tag not in ancestors:
                style_d[style] = {} if tag == "span" else False
    return style_d


def process_p_child_tags(
    doc: Document,
    paragraph: Paragraph,
    tag: Tag,
    parent_tag: Tag,
    cell: _Cell,
    styles=None,
) -> None:
    """Processes child tags of a p tag or other tags like <td>, <li>..."""
    for child in tag.children:
        ancestors = [tag.name] + [t.name for t in tag.parents]
        if (
            isinstance(child, str)
            and (child.parent.name == tag.name)
            and (not skip_crlf(child))
        ):
            if not styles:
                run = paragraph.add_run()
                run.text = child
            else:
                styles = check_style_parent(
                    ancestors,
                    styles,
                    {
                        "span": "span",
                        ("b", "strong"): "bold",
                        ("i", "em"): "italic",
                        "u": "underline",
                    },
                )
                add_text_color(paragraph, child, styles)

        # add new line or line-break
        elif child.name == "br":
            run = paragraph.add_run()
            run.add_break()

        # add <p> styles and child tags
        elif child.name == "p":
            if "id" in child.attrs:
                create_bookmark_run(paragraph, child["id"], "", child["id"])
            if cell is not None:
                process_p_child_tags(
                    doc, cell.add_paragraph(), child, tag, cell, styles
                )
            else:
                process_p_child_tags(doc, paragraph, child, tag, cell, styles)

        # add headings
        elif child.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            lvl = int(child.name[1])
            heading = cell.add_paragraph("", style=f"Heading {lvl}")
            process_p_child_tags(doc, heading, child, parent_tag, cell)

        # add <span> styles and child tags
        elif child.name == "span":
            span_styles = parse_styles(child.get("style", ""))
            styles.update({"span": span_styles})
            if "class" in child.attrs and "bookmark" in child["class"]:
                create_bookmark_run(paragraph, child["name"], child.text, child["id"])
            elif "class" in child.attrs and "anchor" in child["class"]:
                create_bookmark_run(paragraph, child["id"], "", child["id"])
            else:
                process_p_child_tags(doc, paragraph, child, tag, cell, styles)

        # add <strong> or <b> styles for bold and child tags
        elif child.name == "strong" or child.name == "b":
            bold_styles = {"bold": True}
            styles.update(bold_styles)
            styles = check_style_parent(ancestors, styles, {"span": "span"})
            process_p_child_tags(doc, paragraph, child, tag, cell, styles)

        # add italic styles <em> or <i>
        elif child.name == "em" or child.name == "i":
            italic_styles = {"italic": True}
            styles.update(italic_styles)
            tags = {
                "span": "span",
                ("b", "strong"): "bold",
            }
            styles = check_style_parent(ancestors, styles, tags)
            process_p_child_tags(doc, paragraph, child, tag, cell, styles)

        # add underline styles <u> and child tags
        elif child.name == "u":
            underline_styles = {"underline": True}
            styles.update(underline_styles)
            tags = {"span": "span", ("b", "strong"): "bold", ("i", "em"): "italic"}
            styles = check_style_parent(ancestors, styles, tags)
            process_p_child_tags(doc, paragraph, child, tag, cell, styles)

        # add mark styles <mark> and child tags
        elif child.name == "mark":
            process_p_child_tags(doc, paragraph, child, tag, cell, styles)

        # add image under <p> tag
        # elif child.name == "img":
        #     img_data = add_images(child)
        #     if isinstance(img_data, str):
        #         skip_image(doc, cell, img_data)
        #     else:
        #         run = paragraph.add_run()
        #         run.add_picture(img_data)
        #         img_data.close()

        # add anchor tags inside <p>
        elif child.name == "a":
            if "class" in child.attrs and "anchor-link" in child["class"]:
                a_id = child["class"][0]
                bookmark_span = soup.find(id=a_id)
                if bookmark_span:
                    bk_name = (
                        bookmark_span["name"]
                        if "name" in bookmark_span.attrs
                        else bookmark_span["id"]
                    )
                    create_internal_hyperlink_run(paragraph, child.text, bk_name)
            else:
                add_links(paragraph, child.text, child.get("href", ""))

        # add lists inside paragraph
        elif child.name == "ul" or child.name == "ol":
            if child.parent.name == "td":
                process_list(cell, doc, child)
            else:
                process_list(None, doc, child)

        # add table inside paragraph
        elif child.name == "table":
            add_docx_tables(doc, child)

        # process blockquote tags
        elif child.name == "blockquote" and child.parent.name in ["td", "th"]:
            process_blockquote_paragraphs(doc, child, cell)


def align_para(style_dict: dict[str, str], paragraph: Paragraph):
    """
    Set Paragraph alignment for docx based on text-align css value
    """
    if "text-align" in style_dict:
        value = style_dict.get("text-align", "left")
        if "center" == value:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "right" == value:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif "justify" == value:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# def add_heading_style(heading: Paragraph, doc: Document, head_tag: Tag):
#     """
#     Add font color to heading and other styles if present
#     """
#     # for child in head_tag.contents:
#     #     if isinstance(child, str):
#     #         heading.add_run(child)
#     #     elif child.name == "span":
#     #         style_head = {"span": parse_styles(child.get("style", ""))}
#     #         add_text_color(heading, child.get_text(), style_head)
#     process_p_child_tags(doc, heading, head_tag, head_tag, None, {})


def set_document_margin(doc: Document) -> None:
    """
    This method sets document margin to the default
    "Normal" in MS-Word
    """
    # set page layout margin to "Normal"
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)


def skip_image(doc: Document, cell: _Cell, err_msg: str) -> None:
    """
    Skip an invalid image and adds a paragraph text with error message
    """
    if cell:
        p = cell.add_paragraph()
        run = p.add_run()
        run.text = err_msg
        run.bold = True
        run.italic = True
    else:
        p = doc.add_paragraph()
        run = p.add_run()
        run.text = err_msg
        run.bold = True
        run.italic = True


def process_blockquote_paragraphs(doc: Document, tag: Tag, cell: _Cell):
    """
    Add <p> tags to document, which are inside <blockquote> tag in html
    """
    for p_tags in tag.children:
        if isinstance(p_tags, str) and (not skip_crlf(p_tags)):
            paragraph = doc.add_paragraph() if cell is None else cell.add_paragraph()
            run = paragraph.add_run()
            run.text = p_tags
        elif not isinstance(p_tags, str):
            paragraph = doc.add_paragraph() if cell is None else cell.add_paragraph()
            process_p_child_tags(doc, paragraph, p_tags, tag, None, {})


def html_to_docx(html: str) -> BytesIO:
    """
    A Utility method to convert html to docx
    Takes an html string as input and returns docx object as bytes
    """
    # Create Docx document object
    doc = Document()

    # Create parsed html object for tree navigation
    global soup
    soup = BeautifulSoup(html, "lxml")

    # Assign title informatio in document properties
    doc.core_properties.title = (
        soup.title.string if soup.title else "Converted Document"
    )

    # Add Document Title to Docx object
    doc.add_heading(doc.core_properties.title, level=0)

    list_prev_p = None
    # Iterate over child tags in document body
    for tag in soup.body.descendants:
        # Add Heading in docx
        if tag.name in ["h1", "h2", "h3", "h4", "h5", "h6"] and tag.parent.name not in [
            "td",
            "th",
        ]:
            lvl = int(tag.name[1])
            heading = doc.add_heading("", level=lvl)
            process_p_child_tags(doc, heading, tag, tag, None, {})
            if "id" in tag.attrs:
                create_bookmark_run(heading, tag["id"], "", tag["id"])
        # Add Pictures
        # elif tag.name == "img" and tag.parent.name not in ["p", "li", "td", "th"]:
        #     img_data = add_images(tag)
        #     if isinstance(img_data, str):
        #         skip_image(doc, None, img_data)
        #     else:
        #         doc.add_picture(img_data)
        #         img_data.close()
        # Skip empty newlines
        elif skip_crlf(tag):
            continue

        # Add Paragraphs and nested tags
        elif tag.name == "p" and tag.parent.name not in [
            "li",
            "td",
            "th",
            "blockquote",
        ]:
            paragraph = doc.add_paragraph()
            style_str = None if "style" not in tag.attrs else tag["style"]
            style_dict = parse_styles(style_str)
            align_para(style_dict, paragraph)
            if "id" in tag.attrs:
                create_bookmark_run(paragraph, tag["id"], "", tag["id"])
            process_p_child_tags(doc, paragraph, tag, tag, None, style_dict)

        # Add paragraphs in <blockquote> tag
        elif tag.name == "blockquote" and tag.parent.name in ["div", "body"]:
            process_blockquote_paragraphs(doc, tag, None)

        # Add div tag content
        elif tag.name == "div":
            if "class" in tag.attrs and "note" in tag["class"]:
                paragraph = doc.add_paragraph()
                process_p_child_tags(doc, paragraph, tag, tag, None, {})

        # Add Lists oustide <p> tags
        elif (tag.name == "ul" or tag.name == "ol") and tag.parent.name not in [
            "td",
            "th",
            "li",
        ]:
            if is_list_continued(tag):
                list_prev_p = process_list(None, doc, tag, parent_paragraph=list_prev_p)
            else:
                list_prev_p = process_list(None, doc, tag)
        # Add HTML Tables to Docx
        elif tag.name == "table" and tag.parent.name not in ["li", "p"]:
            add_docx_tables(doc, tag)
        # Add Hyperlinks [or anchor links based on structure]
        elif tag.name == "a" and tag.parent.name not in [
            "p",
            "li",
            "td",
            "th",
            "h1",
            "h2",
            "h3",
            "h4",
            "h5",
            "h6",
        ]:
            add_links(doc.add_paragraph(), tag.text, tag["href"])

    set_document_margin(doc)
    # return docx output
    io = BytesIO()
    doc.save(io)
    io.seek(0)
    return io
